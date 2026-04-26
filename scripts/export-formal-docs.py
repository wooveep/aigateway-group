#!/usr/bin/env python3
from __future__ import annotations

import argparse
import html
import re
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:
    print(
        "Missing Python dependency: python-docx.\n"
        "Recommended bootstrap:\n"
        "  python3 -m venv .venv-docs\n"
        "  .venv-docs/bin/pip install python-docx\n"
        "  .venv-docs/bin/python scripts/export-formal-docs.py",
        file=sys.stderr,
    )
    raise SystemExit(1) from exc


REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_VERSION = "1.1.0"


def default_docs(version: str) -> list[Path]:
    candidates = [
        REPO_ROOT / f"docs/overview/aigateway-whitepaper-{version}.md",
        REPO_ROOT / f"docs/release/{version}/release-notes.md",
        REPO_ROOT / f"docs/release/{version}/image-bundle.md",
        REPO_ROOT / f"docs/release/{version}/deployment-guide.md",
        REPO_ROOT / f"docs/manual/{version}/user-manual.md",
    ]
    return [path for path in candidates if path.exists()]


@dataclass
class Block:
    kind: str
    value: object
    extra: str = ""


def parse_markdown(markdown_text: str) -> list[Block]:
    lines = markdown_text.splitlines()
    blocks: list[Block] = []
    paragraph: list[str] = []
    list_items: list[str] = []
    list_kind = ""
    code_lines: list[str] = []
    code_lang = ""
    in_code = False

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append(Block("paragraph", " ".join(line.strip() for line in paragraph).strip()))
            paragraph = []

    def flush_list() -> None:
        nonlocal list_items, list_kind
        if list_items:
            blocks.append(Block(list_kind, list_items[:]))
            list_items = []
            list_kind = ""

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                blocks.append(Block("code", "\n".join(code_lines), code_lang))
                code_lines = []
                code_lang = ""
                in_code = False
            else:
                flush_paragraph()
                flush_list()
                in_code = True
                code_lang = stripped[3:].strip()
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not stripped:
            flush_paragraph()
            flush_list()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            flush_list()
            blocks.append(Block("heading", heading_match.group(2).strip(), heading_match.group(1)))
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if image_match:
            flush_paragraph()
            flush_list()
            blocks.append(Block("image", image_match.group(2).strip(), image_match.group(1).strip()))
            continue

        bullet_match = re.match(r"^- (.*)$", stripped)
        if bullet_match:
            flush_paragraph()
            if list_kind not in ("", "bullet"):
                flush_list()
            list_kind = "bullet"
            list_items.append(bullet_match.group(1).strip())
            continue

        number_match = re.match(r"^\d+\. (.*)$", stripped)
        if number_match:
            flush_paragraph()
            if list_kind not in ("", "number"):
                flush_list()
            list_kind = "number"
            list_items.append(number_match.group(1).strip())
            continue

        paragraph.append(stripped)

    if in_code:
        blocks.append(Block("code", "\n".join(code_lines), code_lang))
    flush_paragraph()
    flush_list()
    return blocks


def parse_inlines(text: str) -> list[tuple[str, str]]:
    tokens: list[tuple[str, str]] = []
    i = 0
    while i < len(text):
        if text.startswith("`", i):
            end = text.find("`", i + 1)
            if end != -1:
                tokens.append(("code", text[i + 1 : end]))
                i = end + 1
                continue
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end != -1:
                tokens.append(("bold", text[i + 2 : end]))
                i = end + 2
                continue
        start = i
        while i < len(text) and not text.startswith("`", i) and not text.startswith("**", i):
            i += 1
        if start != i:
            tokens.append(("text", text[start:i]))
    return tokens


def add_inline_runs(paragraph, text: str) -> None:
    for kind, value in parse_inlines(text):
        run = paragraph.add_run(value)
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(9)


def render_docx(doc_path: Path, output_path: Path) -> None:
    source_text = doc_path.read_text(encoding="utf-8")
    blocks = parse_markdown(source_text)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.start_type = WD_SECTION.NEW_PAGE

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    code_style = styles["No Spacing"]
    code_style.font.name = "Courier New"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code_style.font.size = Pt(9)

    for block in blocks:
        if block.kind == "heading":
            level = min(len(str(block.extra)), 4)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(paragraph, str(block.value))
            continue

        if block.kind == "paragraph":
            paragraph = doc.add_paragraph()
            add_inline_runs(paragraph, str(block.value))
            continue

        if block.kind == "bullet":
            for item in block.value:
                paragraph = doc.add_paragraph(style="List Bullet")
                add_inline_runs(paragraph, str(item))
            continue

        if block.kind == "number":
            for item in block.value:
                paragraph = doc.add_paragraph(style="List Number")
                add_inline_runs(paragraph, str(item))
            continue

        if block.kind == "code":
            for line in str(block.value).splitlines() or [""]:
                paragraph = doc.add_paragraph(style="No Spacing")
                run = paragraph.add_run(line)
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
                run.font.size = Pt(9)
            continue

        if block.kind == "image":
            image_path = (doc_path.parent / str(block.value)).resolve()
            if image_path.exists():
                if block.extra:
                    doc.add_paragraph(str(block.extra), style="Caption")
                doc.add_picture(str(image_path), width=Cm(16.5))
            else:
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"[Missing image] {image_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def inline_html(text: str) -> str:
    parts: list[str] = []
    for kind, value in parse_inlines(text):
        escaped = html.escape(value)
        if kind == "bold":
            parts.append(f"<strong>{escaped}</strong>")
        elif kind == "code":
            parts.append(f"<code>{escaped}</code>")
        else:
            parts.append(escaped)
    return "".join(parts)


def render_html(doc_path: Path, output_path: Path) -> None:
    source_text = doc_path.read_text(encoding="utf-8")
    blocks = parse_markdown(source_text)
    body_parts: list[str] = []

    for block in blocks:
        if block.kind == "heading":
            level = min(len(str(block.extra)), 6)
            body_parts.append(f"<h{level}>{inline_html(str(block.value))}</h{level}>")
        elif block.kind == "paragraph":
            body_parts.append(f"<p>{inline_html(str(block.value))}</p>")
        elif block.kind == "bullet":
            items = "".join(f"<li>{inline_html(str(item))}</li>" for item in block.value)
            body_parts.append(f"<ul>{items}</ul>")
        elif block.kind == "number":
            items = "".join(f"<li>{inline_html(str(item))}</li>" for item in block.value)
            body_parts.append(f"<ol>{items}</ol>")
        elif block.kind == "code":
            body_parts.append(f"<pre><code>{html.escape(str(block.value))}</code></pre>")
        elif block.kind == "image":
            image_path = (doc_path.parent / str(block.value)).resolve()
            caption = f"<figcaption>{html.escape(block.extra)}</figcaption>" if block.extra else ""
            if image_path.exists():
                uri = image_path.as_uri()
                body_parts.append(
                    f"<figure><img src=\"{uri}\" alt=\"{html.escape(block.extra)}\">{caption}</figure>"
                )
            else:
                body_parts.append(f"<p>[Missing image] {html.escape(str(image_path))}</p>")

    document_title = html.escape(doc_path.stem)
    html_text = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{document_title}</title>
  <style>
    @page {{
      size: A4;
      margin: 16mm;
    }}
    body {{
      color: #1f2937;
      font-family: "Noto Sans SC", "Microsoft YaHei", Arial, sans-serif;
      font-size: 12px;
      line-height: 1.65;
      margin: 0 auto;
      max-width: 180mm;
    }}
    h1, h2, h3, h4 {{
      color: #0f172a;
      line-height: 1.3;
      margin: 1.1em 0 0.4em;
    }}
    h1 {{
      border-bottom: 2px solid #dbe4ee;
      padding-bottom: 0.25em;
    }}
    p, ul, ol, pre, figure {{
      margin: 0.55em 0;
    }}
    ul, ol {{
      padding-left: 1.5em;
    }}
    pre {{
      background: #f6f8fa;
      border: 1px solid #d0d7de;
      border-radius: 6px;
      overflow-x: auto;
      padding: 12px;
      white-space: pre-wrap;
    }}
    code {{
      background: #eff3f6;
      border-radius: 4px;
      font-family: "SFMono-Regular", Consolas, "Liberation Mono", monospace;
      font-size: 0.92em;
      padding: 0.12em 0.32em;
    }}
    pre code {{
      background: transparent;
      padding: 0;
    }}
    img {{
      border: 1px solid #d0d7de;
      border-radius: 6px;
      display: block;
      height: auto;
      max-width: 100%;
    }}
    figcaption {{
      color: #475569;
      font-size: 11px;
      margin-top: 0.4em;
      text-align: center;
    }}
  </style>
</head>
<body>
{''.join(body_parts)}
</body>
</html>
"""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(html_text, encoding="utf-8")


def render_pdf(html_path: Path, pdf_path: Path, chrome_bin: str) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            chrome_bin,
            "--headless=new",
            "--disable-gpu",
            "--allow-file-access-from-files",
            "--print-to-pdf-no-header",
            f"--print-to-pdf={pdf_path}",
            html_path.as_uri(),
        ],
        check=True,
        cwd=REPO_ROOT,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def build_manifest(rows: Iterable[tuple[Path, Path, Path, Path]], output_path: Path, version: str) -> None:
    lines = [
        f"# AIGateway {version} 正式文档导出清单",
        "",
        "本目录存放从 `docs/` Markdown 真相源导出的正式交付件。",
        "",
    ]
    for source_md, html_path, docx_path, pdf_path in rows:
        lines.extend(
            [
                f"## {source_md.name}",
                "",
                f"- Source: `{source_md.relative_to(REPO_ROOT)}`",
                f"- HTML: `{html_path.relative_to(REPO_ROOT)}`",
                f"- DOCX: `{docx_path.relative_to(REPO_ROOT)}`",
                f"- PDF: `{pdf_path.relative_to(REPO_ROOT)}`",
                "",
            ]
        )
    output_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Export formal Markdown docs to DOCX and PDF.")
    parser.add_argument(
        "--version",
        default=DEFAULT_VERSION,
        help=f"Release documentation version. Default: {DEFAULT_VERSION}.",
    )
    parser.add_argument(
        "--output-dir",
        default="",
        help="Directory for exported HTML, DOCX, and PDF files.",
    )
    parser.add_argument(
        "--chrome-bin",
        default="/usr/bin/google-chrome",
        help="Chrome or Chromium binary used for PDF rendering.",
    )
    args = parser.parse_args()

    output_dir = Path(args.output_dir or REPO_ROOT / f"out/docs/{args.version}").resolve()
    html_dir = output_dir / "html"
    rows: list[tuple[Path, Path, Path, Path]] = []
    docs = default_docs(args.version)
    if not docs:
        raise SystemExit(f"No formal docs found for version {args.version}")

    for doc_path in docs:
        stem = doc_path.stem
        html_path = html_dir / f"{stem}.html"
        docx_path = output_dir / f"{stem}.docx"
        pdf_path = output_dir / f"{stem}.pdf"

        render_html(doc_path, html_path)
        render_docx(doc_path, docx_path)
        render_pdf(html_path, pdf_path, args.chrome_bin)

        rows.append((doc_path, html_path, docx_path, pdf_path))

    build_manifest(rows, output_dir / "README.md", args.version)
    print(f"Exported {len(rows)} documents to {output_dir}")


if __name__ == "__main__":
    main()
